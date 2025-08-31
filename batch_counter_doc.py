import zipfile
from os import scandir
from docx import Document

ru_chars = 'абвгдеёжзийклмнопрстуфхцчшщыэюя'
rate = 1.3


def count_ru(file):
    ru = []
    document = Document(file)
    # print(len(document.paragraphs))
    for par in document.paragraphs:
        for run in par.runs:
            # print(run.text)
            for word in run.text.split():
                for char in word.lower():
                    if char in ru_chars:
                        # print(word)
                        ru.append(word)
                        break
    return len(ru)


def get_files_ext(path, extension):
    files = [item.path for item in scandir(path) if item.name.endswith(extension)]
    return files


def dirs_get_files_filter(path, string):
    files = []
    dirs = [item.path for item in scandir(path) if item.is_dir()]
    for d in dirs:
        files += [item.path for item in scandir(d) if string.lower() not in item.path.lower()]
    return files


def zip_batch_count(path, extension='zip'):  # doesn't work with russian filenames
    count = 0
    for file in get_files_ext(path, extension):
        print(f'processing {file}...')
        with zipfile.ZipFile(file) as arch:
            print(arch.namelist())
            for name in arch.namelist():
                print(' ', name, end=': ')
                if name.endswith('docx'):
                    # with arch.open(name) as cur:
                    count += count_ru(arch.open(name))
                    print(count)
    return count


def dirs_batch_count(path, string):
    return sum(count_ru(file) for file in dirs_get_files_filter(path, string))


# file = r'D:\Earn\Translate\English\Atelier\VideoTrans\atelier-mr.ru_2020-09-18_21-55\Марина Сипович.docx'
# print(count_ru(file))

path = r'D:\Earn\Translate\English\Atelier\VideoTrans\Batch2'
string = 'eng.'

total_words = dirs_batch_count(path, string)
print(total_words, total_words * rate)
